"""Document Library router (PRD §6 / Sprint 3 stories S3-001 → S3-010).

Architecture:
  Document        — parent record (title, type, status, ownership, timestamps)
  DocumentVersion — immutable content snapshots; current_version_id on Document
                    points at the latest
  DocumentTag     — many tags per document (S3-006 filter)
  JobDocumentLink — N:N between a Job and a specific DocumentVersion (S3-009)

File-storage helpers (PDF/DOCX/text extraction + write-back) live in this module
so the upload + edit workflow keeps its existing semantics.
"""

from __future__ import annotations

import base64
import os
from datetime import date as date_class

from docx import Document as DocxDocument
from fastapi import APIRouter, Depends, File, Form, HTTPException, UploadFile, status
from PyPDF2 import PdfReader
from sqlalchemy.orm import Session

from database import get_db
from database.auth import get_current_user
from database.database import get_settings
from database.models.document import (
    create_document,
    get_document,
    get_documents_for_user,
    hard_delete_document,
    update_document,
)
from database.models.document_tag import (
    add_tag,
    get_tags_for_document,
    remove_tag,
)
from database.models.document_version import (
    create_document_version,
    get_document_version,
    get_versions_for_document,
)
from database.models.education import get_educations_by_user
from database.models.experience import get_experiences_by_user
from database.models.job import get_job
from database.models.job_document_link import (
    get_links_for_job,
    link_version_to_job,
    unlink,
)
from database.models.profile import get_profile_by_user_id
from database.models.skill import get_skills_for_user
from database.models.user import User
from schemas import (
    DocumentCreate,
    DocumentResponse,
    DocumentTagCreate,
    DocumentTagResponse,
    DocumentUpdate,
    DocumentVersionCreate,
    DocumentVersionResponse,
    JobDocumentLinkCreate,
    JobDocumentLinkResponse,
)

router = APIRouter()

_ROUTERS_DIR = os.path.dirname(os.path.abspath(__file__))
_BACKEND_DIR = os.path.dirname(_ROUTERS_DIR)
UPLOAD_BASE = os.path.join(_BACKEND_DIR, "uploads")


# --------------------------------------------------------------------------- #
#  File-storage helpers                                                         #
# --------------------------------------------------------------------------- #


def _extract_pdf_content(file_path: str) -> str:
    with open(file_path, "rb") as f:
        pdf_reader = PdfReader(f)
        return "\n".join(p.extract_text() or "" for p in pdf_reader.pages).strip()


def _extract_docx_content(file_path: str) -> str:
    doc = DocxDocument(file_path)
    return "\n".join(p.text for p in doc.paragraphs).strip()


def _write_docx_content(file_path: str, content: str) -> None:
    doc = DocxDocument()
    for line in content.split("\n"):
        doc.add_paragraph(line)
    doc.save(file_path)


def _write_pdf_content(file_path: str, content: str) -> None:
    from io import BytesIO

    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72,
    )
    story = []
    styles = getSampleStyleSheet()
    for line in content.split("\n"):
        if line.strip():
            story.append(Paragraph(line, styles["Normal"]))
            story.append(Spacer(1, 0.12 * inch))
    doc.build(story)
    buffer.seek(0)
    with open(file_path, "wb") as f:
        f.write(buffer.read())


def _update_file_content(file_path: str, filename: str, content: str) -> None:
    ext = os.path.splitext(filename)[1].lower()
    if ext in (".txt", ".md"):
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(content)
    elif ext == ".docx":
        _write_docx_content(file_path, content)
    elif ext == ".pdf":
        _write_pdf_content(file_path, content)
    else:
        raise ValueError(f"Unsupported file type for editing: {ext}")


def _read_file(file_path: str, filename: str) -> dict:
    ext = os.path.splitext(filename or "")[1].lower()
    if ext == ".pdf":
        text = _extract_pdf_content(file_path)
        with open(file_path, "rb") as f:
            b64 = base64.b64encode(f.read()).decode("utf-8")
        return {"content": text, "format": "pdf", "binary_data": b64, "editable": True}
    if ext == ".docx":
        return {
            "content": _extract_docx_content(file_path),
            "format": "docx",
            "editable": True,
        }
    if ext in (".txt", ".md"):
        with open(file_path, "r", encoding="utf-8") as f:
            return {"content": f.read(), "format": ext[1:], "editable": True}
    return {
        "content": f"[Unsupported file type: {ext}]",
        "format": "unknown",
        "editable": False,
    }


def _build_upload_path(
    base: str, first_name: str, last_name: str, user_id: int, filename: str
) -> str:
    last_initial = (last_name or "X")[0].upper()
    first_initial = (first_name or "X")[0].upper()
    full_name = f"{first_name or ''} {last_name or ''}".strip() or f"user_{user_id}"
    return os.path.join(
        base, last_initial, first_initial, full_name, str(user_id), filename
    )


def _ensure_owns(document, current_user: User) -> None:
    if document is None:
        raise HTTPException(status_code=404, detail="Document not found")
    if document.user_id != current_user.user_id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN,
            detail="You do not have permission to access this document",
        )


def _build_job_context(job) -> str:
    parts = [
        "\nTARGET JOB:",
        f"Title: {job.title}",
        f"Company: {job.company_name}",
    ]
    if job.location:
        parts.append(f"Location: {job.location}")
    if job.salary:
        parts.append(f"Salary: {job.salary}")
    if job.description:
        parts.append(f"Job Description:\n{job.description}")
    if job.years_of_experience is not None:
        parts.append(f"Years of Experience Expected: {job.years_of_experience}")
    if job.company_research_notes:
        parts.append(f"Company Research:\n{job.company_research_notes}")
    return "\n".join(parts)


# --------------------------------------------------------------------------- #
#  Document CRUD (list, read, create, update, archive/restore, hard delete)     #
# --------------------------------------------------------------------------- #


@router.get("/me", response_model=list[DocumentResponse])
def list_my_documents(
    include_archived: bool = False,
    document_type: str | None = None,
    status_filter: str | None = None,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """List the user's documents. Supports S3-006 filtering by type and status."""
    docs = get_documents_for_user(
        session, current_user.user_id, include_deleted=include_archived
    )
    if document_type:
        docs = [d for d in docs if d.document_type == document_type]
    if status_filter:
        docs = [d for d in docs if d.status == status_filter]
    return docs


@router.post("", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
def create_document_endpoint(
    body: DocumentCreate,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Create a Document plus an initial v1 DocumentVersion in one shot."""
    doc = create_document(
        session,
        current_user.user_id,
        body.title,
        body.document_type,
        status=body.status or "Draft",
    )
    if body.content or body.storage_location:
        version = create_document_version(
            session,
            doc.document_id,
            storage_location=body.storage_location,
            content=body.content,
            source=body.source or "manual",
        )
        update_document(session, doc.document_id, current_version_id=version.version_id)
    if body.tags:
        for label in body.tags:
            add_tag(session, doc.document_id, label)
    if body.job_id is not None:
        job = get_job(session, body.job_id)
        if job and job.user_id == current_user.user_id:
            current_version_id = doc.current_version_id or (
                create_document_version(
                    session, doc.document_id, source=body.source or "manual"
                ).version_id
            )
            link_version_to_job(
                session,
                job_id=body.job_id,
                version_id=current_version_id,
                role=body.role or doc.document_type,
            )
    return get_document(session, doc.document_id)


@router.get("/{document_id}", response_model=DocumentResponse)
def read_document(
    document_id: int,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = get_document(session, document_id)
    _ensure_owns(doc, current_user)
    return doc


@router.put("/{document_id}", response_model=DocumentResponse)
def update_document_endpoint(
    document_id: int,
    body: DocumentUpdate,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Update document metadata only (S3-002, S3-006, S3-008).

    Use POST /{document_id}/versions to write new content.
    """
    doc = get_document(session, document_id)
    _ensure_owns(doc, current_user)
    return update_document(
        session,
        document_id,
        title=body.title,
        document_type=body.document_type,
        status=body.status,
        is_deleted=body.is_deleted,
    )


@router.delete("/{document_id}", status_code=status.HTTP_204_NO_CONTENT)
def hard_delete_document_endpoint(
    document_id: int,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Hard delete (cascades to versions/tags/links). Prefer PUT with is_deleted=True for archive."""
    doc = get_document(session, document_id)
    _ensure_owns(doc, current_user)
    # Best-effort: remove uploaded file if any version stored a file location
    for version in get_versions_for_document(session, document_id):
        if version.storage_location and os.path.exists(version.storage_location):
            try:
                os.remove(version.storage_location)
            except Exception:
                pass
    hard_delete_document(session, document_id)


@router.post(
    "/{document_id}/duplicate",
    response_model=DocumentResponse,
    status_code=status.HTTP_201_CREATED,
)
def duplicate_document(
    document_id: int,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """S3-007: duplicate an existing document (copies the current version content)."""
    src = get_document(session, document_id)
    _ensure_owns(src, current_user)
    new_doc = create_document(
        session,
        current_user.user_id,
        f"{src.title} (copy)",
        src.document_type,
        status=src.status,
    )
    # Copy the current version's content forward as v1 of the new document.
    if src.current_version_id:
        cv = get_document_version(session, src.current_version_id)
        if cv:
            new_version = create_document_version(
                session,
                new_doc.document_id,
                storage_location=None,  # new doc gets its own file lifecycle
                content=cv.content,
                source="duplicate",
            )
            update_document(
                session, new_doc.document_id, current_version_id=new_version.version_id
            )
    return get_document(session, new_doc.document_id)


# --------------------------------------------------------------------------- #
#  Versions                                                                     #
# --------------------------------------------------------------------------- #


@router.get("/{document_id}/versions", response_model=list[DocumentVersionResponse])
def list_versions(
    document_id: int,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = get_document(session, document_id)
    _ensure_owns(doc, current_user)
    return get_versions_for_document(session, document_id)


@router.post(
    "/{document_id}/versions",
    response_model=DocumentVersionResponse,
    status_code=status.HTTP_201_CREATED,
)
def create_new_version(
    document_id: int,
    body: DocumentVersionCreate,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Append a new version. Sets it as the document's current_version."""
    doc = get_document(session, document_id)
    _ensure_owns(doc, current_user)
    version = create_document_version(
        session,
        document_id,
        storage_location=body.storage_location,
        content=body.content,
        source=body.source or "manual",
    )
    update_document(session, document_id, current_version_id=version.version_id)
    return version


# --------------------------------------------------------------------------- #
#  Content read / edit (operates on the current version)                        #
# --------------------------------------------------------------------------- #


@router.get("/{document_id}/content")
def read_current_content(
    document_id: int,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = get_document(session, document_id)
    _ensure_owns(doc, current_user)
    if doc.current_version_id is None:
        raise HTTPException(status_code=404, detail="Document has no version yet")
    version = get_document_version(session, doc.current_version_id)
    if version is None:
        raise HTTPException(status_code=404, detail="Current version missing")
    if version.content:
        return {"content": version.content, "format": "text"}
    if version.storage_location and os.path.exists(version.storage_location):
        try:
            return _read_file(version.storage_location, doc.title)
        except Exception as e:  # noqa: BLE001
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to read file: {e}",
            )
    raise HTTPException(status_code=404, detail="No content available")


@router.put("/{document_id}/content", response_model=DocumentVersionResponse)
def edit_current_content(
    document_id: int,
    body: dict,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Edit current version content. If the version is file-backed, the file is
    rewritten on disk; otherwise a new version row is appended carrying the
    text content.
    """
    doc = get_document(session, document_id)
    _ensure_owns(doc, current_user)
    content = body.get("content")
    if content is None:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Content field is required",
        )
    current = (
        get_document_version(session, doc.current_version_id)
        if doc.current_version_id
        else None
    )
    if (
        current
        and current.storage_location
        and os.path.exists(current.storage_location)
    ):
        try:
            _update_file_content(
                current.storage_location,
                os.path.basename(current.storage_location),
                content,
            )
        except ValueError as e:
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail=str(e))
        # Append a new version row with the in-memory text snapshot for history.
        new_version = create_document_version(
            session,
            document_id,
            storage_location=current.storage_location,
            content=content,
            source="edit",
        )
    else:
        new_version = create_document_version(
            session, document_id, content=content, source="edit"
        )
    update_document(session, document_id, current_version_id=new_version.version_id)
    return new_version


# --------------------------------------------------------------------------- #
#  Upload (S3-004)                                                              #
# --------------------------------------------------------------------------- #


@router.post(
    "/upload", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED
)
async def upload_document(
    file: UploadFile = File(...),
    document_type: str = Form(...),
    title: str | None = Form(None),
    status_value: str = Form("Draft"),
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    profile = get_profile_by_user_id(session, current_user.user_id)
    if not profile:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="User profile not found — create a profile before uploading documents",
        )
    dest_path = _build_upload_path(
        UPLOAD_BASE,
        profile.first_name,
        profile.last_name,
        current_user.user_id,
        file.filename,
    )
    os.makedirs(os.path.dirname(dest_path), exist_ok=True)
    contents = await file.read()
    with open(dest_path, "wb") as f:
        f.write(contents)

    doc = create_document(
        session,
        current_user.user_id,
        title or file.filename,
        document_type,
        status=status_value,
    )
    version = create_document_version(
        session,
        doc.document_id,
        storage_location=dest_path,
        source="upload",
    )
    update_document(session, doc.document_id, current_version_id=version.version_id)
    return get_document(session, doc.document_id)


# --------------------------------------------------------------------------- #
#  Tags (S3-006)                                                                #
# --------------------------------------------------------------------------- #


@router.get("/{document_id}/tags", response_model=list[DocumentTagResponse])
def list_tags(
    document_id: int,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = get_document(session, document_id)
    _ensure_owns(doc, current_user)
    return get_tags_for_document(session, document_id)


@router.post(
    "/{document_id}/tags",
    response_model=DocumentTagResponse,
    status_code=status.HTTP_201_CREATED,
)
def create_tag(
    document_id: int,
    body: DocumentTagCreate,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = get_document(session, document_id)
    _ensure_owns(doc, current_user)
    tag = add_tag(session, document_id, body.label)
    if tag is None:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST, detail="Empty tag label"
        )
    return tag


@router.delete("/{document_id}/tags/{label}", status_code=status.HTTP_204_NO_CONTENT)
def delete_tag(
    document_id: int,
    label: str,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    doc = get_document(session, document_id)
    _ensure_owns(doc, current_user)
    if not remove_tag(session, document_id, label):
        raise HTTPException(status_code=404, detail="Tag not found")


# --------------------------------------------------------------------------- #
#  Job ↔ Document linking (S3-009 / S3-010)                                     #
# --------------------------------------------------------------------------- #


@router.post(
    "/links",
    response_model=JobDocumentLinkResponse,
    status_code=status.HTTP_201_CREATED,
)
def create_link(
    body: JobDocumentLinkCreate,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    job = get_job(session, body.job_id)
    if not job or job.user_id != current_user.user_id:
        raise HTTPException(status_code=404, detail="Job not found")
    version = get_document_version(session, body.version_id)
    if not version:
        raise HTTPException(status_code=404, detail="Version not found")
    parent_doc = get_document(session, version.document_id)
    _ensure_owns(parent_doc, current_user)
    return link_version_to_job(
        session, job_id=body.job_id, version_id=body.version_id, role=body.role
    )


@router.delete("/links/{link_id}", status_code=status.HTTP_204_NO_CONTENT)
def delete_link(
    link_id: int,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    # Ownership inferred via the underlying job.
    from database.models.job_document_link import JobDocumentLink

    link = session.get(JobDocumentLink, link_id)
    if link is None:
        raise HTTPException(status_code=404, detail="Link not found")
    job = get_job(session, link.job_id)
    if not job or job.user_id != current_user.user_id:
        raise HTTPException(
            status_code=status.HTTP_403_FORBIDDEN, detail="Access denied"
        )
    unlink(session, link_id)


@router.get(
    "/links/by-job/{job_id}",
    response_model=list[JobDocumentLinkResponse],
)
def list_links_for_job(
    job_id: int,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    job = get_job(session, job_id)
    if not job or job.user_id != current_user.user_id:
        raise HTTPException(status_code=404, detail="Job not found")
    return get_links_for_job(session, job_id)


# --------------------------------------------------------------------------- #
#  AI generation (creates a Document + version + optional job link)             #
# --------------------------------------------------------------------------- #


def _build_profile_text(session: Session, user_id: int) -> str:
    profile = get_profile_by_user_id(session, user_id)
    if not profile:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="User profile not found — create a profile first",
        )
    experiences = get_experiences_by_user(session, user_id)
    educations = get_educations_by_user(session, user_id)
    skills = get_skills_for_user(session, user_id)

    lines = [f"Name: {profile.first_name} {profile.last_name}"]
    if profile.phone_number:
        lines.append(f"Phone: {profile.phone_number}")
    if profile.summary:
        lines.append(f"Summary: {profile.summary}")
    if experiences:
        lines.append("\nWORK EXPERIENCE:")
        for exp in experiences:
            end = exp.end_date.strftime("%b %Y") if exp.end_date else "Present"
            lines.append(
                f"  {exp.title} at {exp.company} ({exp.start_date.strftime('%b %Y')} - {end})"
            )
            if exp.location:
                lines.append(f"  Location: {exp.location}")
            if exp.description:
                lines.append(f"  {exp.description}")
    if educations:
        lines.append("\nEDUCATION:")
        for edu in educations:
            field = f" in {edu.field_of_study}" if edu.field_of_study else ""
            lines.append(f"  {edu.degree}{field} — {edu.school}")
            if edu.gpa:
                lines.append(f"  GPA: {edu.gpa}")
            if edu.start_date and edu.end_date:
                lines.append(f"  {edu.start_date.year} - {edu.end_date.year}")
    if skills:
        lines.append("\nSKILLS:")
        lines.append("  " + ", ".join(s.name for s in skills))
    return "\n".join(lines)


def _call_openai(system_prompt: str, user_message: str) -> str:
    import openai

    settings = get_settings()
    api_key = os.environ.get("OPENAI_API_KEY") or settings.openai_api_key
    if not api_key:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="OpenAI API key is not configured.",
        )
    try:
        client = openai.OpenAI(api_key=api_key)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_message},
            ],
            max_tokens=4096,
            temperature=0.7,
        )
        return response.choices[0].message.content.strip()
    except openai.AuthenticationError:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Invalid OpenAI API key.",
        )
    except openai.RateLimitError:
        raise HTTPException(
            status_code=status.HTTP_429_TOO_MANY_REQUESTS,
            detail="OpenAI rate limit reached. Try again shortly.",
        )
    except Exception as e:  # noqa: BLE001
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"AI generation failed: {e}",
        )


def _generate_doc(
    session: Session,
    *,
    user: User,
    document_type: str,
    role: str,
    title_prefix: str,
    system_prompt: str,
    body: dict,
) -> dict:
    profile_text = _build_profile_text(session, user.user_id)
    job_id = body.get("job_id")
    job_context = ""
    job_label = ""
    job = None
    if job_id:
        job = get_job(session, job_id)
        if not job or job.user_id != user.user_id:
            raise HTTPException(status_code=404, detail="Job not found")
        job_context = _build_job_context(job)
        job_label = f" - {job.title} @ {job.company_name}"

    instructions = (body.get("instructions") or "").strip()
    user_message = f"Generate from this profile:\n\n{profile_text}"
    if job_context:
        user_message += f"\n\n{job_context}"
    if instructions:
        user_message += f"\n\nAdditional instructions: {instructions}"

    content = _call_openai(system_prompt, user_message)

    doc_title = f"{title_prefix}{job_label} - {date_class.today().isoformat()}"
    doc = create_document(
        session, user.user_id, doc_title, document_type, status="Draft"
    )
    version = create_document_version(
        session, doc.document_id, content=content, source="ai"
    )
    update_document(session, doc.document_id, current_version_id=version.version_id)
    if job is not None:
        link_version_to_job(
            session, job_id=job.job_id, version_id=version.version_id, role=role
        )
    return {
        "document_id": doc.document_id,
        "version_id": version.version_id,
        "content": content,
        "title": doc_title,
    }


@router.post("/generate-resume")
def generate_resume(
    body: dict,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    return _generate_doc(
        session,
        user=current_user,
        document_type="Resume",
        role="resume",
        title_prefix="AI Resume",
        system_prompt=(
            "You are an expert resume writer. Produce a clean ATS-friendly resume from the "
            "candidate profile and (when given) the target job. Plain text. Standard sections "
            "(CONTACT, SUMMARY, EXPERIENCE, EDUCATION, SKILLS). No code fences or preamble."
        ),
        body=body,
    )


@router.post("/generate-cover-letter")
def generate_cover_letter(
    body: dict,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    return _generate_doc(
        session,
        user=current_user,
        document_type="Cover Letter",
        role="cover_letter",
        title_prefix="AI Cover Letter",
        system_prompt=(
            "You are an expert cover-letter writer. Produce a personalized 3-4 paragraph "
            "cover letter for the target job using the candidate profile. Address it to the "
            "hiring team at the company. End with a clear call to action and the candidate's "
            "name. Plain text only — no preamble or markdown."
        ),
        body=body,
    )


# --------------------------------------------------------------------------- #
#  Backwards-compatible aliases (DocumentCreate shape used by older callers)    #
# --------------------------------------------------------------------------- #


@router.post("/", response_model=DocumentResponse, status_code=status.HTTP_201_CREATED)
def create_document_legacy_path(
    body: DocumentCreate,
    session: Session = Depends(get_db),
    current_user: User = Depends(get_current_user),
):
    """Trailing-slash alias of POST /. Retained for the existing frontend client."""
    return create_document_endpoint(body, session, current_user)
